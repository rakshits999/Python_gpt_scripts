import csv
import openai
from docx import Document
import os
import time

openai.api_key = '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'

progress_file = "progress.txt"

try:
    with open(progress_file, "r") as f:
        progress = f.read().splitlines()
        last_processed_row = int(progress[0])
except FileNotFoundError:
    last_processed_row = 0
    

with open('finalcsv.csv', 'r', encoding='latin-1') as file:
    csvReader = csv.reader(file, delimiter=';')

    doc = Document()
    doc_filename = None
    combined_doc = Document() 

    for j, row in enumerate(csvReader):
        if j < last_processed_row:
            continue

        Topic = None
        prompt = None
        conversation = []
        table_content = None  

        if j in range(30,31):
            for i, prompts in enumerate(row, start=1):
                if i == 3:
                    Topic = prompts
                elif i >= 4:
                    if i == 4:
                        print(f"Column {i}: {prompts}") 
                    else:
                        print(f"Value {i}: {prompts}")

                    doc_filename = " ".join(Topic.split()[:5]) + ".docx"
                    user_message = {"role": "user", "content": prompts}
                    conversation.append(user_message)
                    print(conversation)

                    response = openai.ChatCompletion.create(
                        model="gpt-3.5-turbo-16k",
                        messages=conversation
                    )
                   
                    
                    chat_response =  response.choices[0].message.content.strip()
                    
                    table_start = chat_response.find("<table>")
                    table_end = chat_response.find("</table>")
                    if table_start != -1 and table_end != -1:
                        table_content = chat_response[table_start + len("<table>"):table_end].strip()

                        # Parse table content
                        rows = table_content.split('</tr>')[:-1]  # Split table rows
                        headers = rows[0].split('</th>')[:-1]  # Split table headers

                        text_before_table = chat_response[:table_start]
                        doc.add_paragraph(text_before_table)
                        combined_doc.add_paragraph(text_before_table)

                        # Remove <tr> tags from the table rows
                        rows = [row.replace('<tr>', '') for row in rows]
                        headers = [header.replace('<tr>', '') for header in headers]

                        # Add table to the combined document
                        doc.add_paragraph()
                        table = doc.add_table(rows=1, cols=len(headers))
                        table.style = 'Table Grid'  # Apply table grid style
                        header_cells = table.rows[0].cells
                        for idx, header in enumerate(headers):
                            header_text = header.replace('<th>', '').strip()
                            if idx < len(header_cells):
                                header_cells[idx].text = header_text

                        for row_data in rows[1:]:
                            data_cells = table.add_row().cells
                            data = row_data.split('</td>')[:-1]  # Split table data
                            for idx, value in enumerate(data):
                                value_text = value.replace('<td>', '').strip()
                                data_cells[idx].text = value_text

                        combined_doc.add_paragraph()
                        table = combined_doc.add_table(rows=1, cols=len(headers))
                        table.style = 'Table Grid'  # Apply table grid style
                        header_cells = table.rows[0].cells
                        for idx, header in enumerate(headers):
                            header_text = header.replace('<th>', '').strip()
                            if idx < len(header_cells):
                                header_cells[idx].text = header_text

                        for row_data in rows[1:]:
                            data_cells = table.add_row().cells
                            data = row_data.split('</td>')[:-1]  # Split table data
                            for idx, value in enumerate(data):
                                value_text = value.replace('<td>', '').strip()
                                data_cells[idx].text = value_text


                        # Add text after table to the combined document
                        text_after_table = chat_response[table_end + len("</table>"):]
                        doc.add_paragraph(text_after_table)
                        combined_doc.add_paragraph(text_after_table)

                    elif '\n\n|' in chat_response:
                        start_index = 0
                        end_index = len(chat_response)

                        if '\n\n|' in chat_response:
                            start_index = chat_response.index('\n\n|') + 1

                        if '|\n\n' in chat_response:
                            end_index = chat_response.rindex('|\n\n') + 1

                        extracted_text_before = chat_response[:start_index].strip()
                        doc.add_paragraph(extracted_text_before)
                        combined_doc.add_paragraph(extracted_text_before)
                      

                        extracted_text = chat_response[start_index:end_index].strip()

                        table_data = extracted_text
                        rows = table_data.strip().split('\n')
                        headers = [header.strip() for header in rows[0].split('|')[1:-1]]
                        data = [[cell.strip() for cell in row.split('|')[1:-1]] for row in rows[2:]]

                        html = ''

                        html += '  <tr>\n'
                        for header in headers:
                            html += f'    <th>{header}</th>\n'
                        html += '  </tr>\n'

                        for row in data:
                            html += '  <tr>\n'
                            for cell in row:
                                html += f'    <td>{cell}</td>\n'
                            html += '  </tr>\n'
                      

                        rows = html.split('</tr>')[:-1]  # Split table rows
                        headers = rows[0].split('</th>')[:-1]

                        rows = [row.replace('<tr>', '') for row in rows]
                        headers = [header.replace('<tr>', '') for header in headers]

                        doc.add_paragraph()
                        table = doc.add_table(rows=1, cols=len(headers))
                        table.style = 'Table Grid'  # Apply table grid style
                        header_cells = table.rows[0].cells
                        for idx, header in enumerate(headers):
                            header_text = header.replace('<th>', '').strip()
                            if idx < len(header_cells):
                                header_cells[idx].text = header_text

                        for row_data in rows[1:]:
                            data_cells = table.add_row().cells
                            data = row_data.split('</td>')[:-1]  # Split table data
                            for idx, value in enumerate(data):
                                value_text = value.replace('<td>', '').strip()
                                data_cells[idx].text = value_text

                        combined_doc.add_paragraph()
                        table = combined_doc.add_table(rows=1, cols=len(headers))
                        table.style = 'Table Grid'  # Apply table grid style
                        header_cells = table.rows[0].cells
                        for idx, header in enumerate(headers):
                            header_text = header.replace('<th>', '').strip()
                            if idx < len(header_cells):
                                header_cells[idx].text = header_text

                        for row_data in rows[1:]:
                            data_cells = table.add_row().cells
                            data = row_data.split('</td>')[:-1]  # Split table data
                            for idx, value in enumerate(data):
                                value_text = value.replace('<td>', '').strip()
                                data_cells[idx].text = value_text

                        extracted_text_after = chat_response[end_index + 1:].strip()
                        doc.add_paragraph(extracted_text_after)
                        combined_doc.add_paragraph(extracted_text_after)
             
                    else:
                        
                        if "<table>" not in chat_response:
                            doc.add_paragraph(chat_response)
                            combined_doc.add_paragraph(chat_response)

                    system_message = {"role": "assistant", "content": chat_response}
                    conversation.append(system_message)



            if doc_filename:
                doc.save(doc_filename)
                doc = Document()

                with open(progress_file, "w") as f:
                    f.write(f"{j + 1}")
                    time.sleep(1)

    combined_doc.save("Combined_Doc.docx")

if last_processed_row == j + 1:
    os.remove(progress_file)
